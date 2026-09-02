from fastapi import APIRouter, HTTPException, Depends, File, UploadFile
from typing import Optional
from datetime import datetime, timezone
import uuid
import logging
import os
import shutil
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
import docx
from middleware.auth_middleware import get_current_user, get_current_user_optional

router = APIRouter(tags=["courses"])
logger = logging.getLogger(__name__)

UPLOAD_DIR = "backend/uploads/courses"
CONTENT_UPLOAD_DIR = "backend/uploads/course_contents"

@router.get("/last-course")
async def get_last_course(current_user: dict = Depends(get_current_user)):
    """Get user's last accessed course"""
    try:
        users_col = get_collection('users')
        user = await users_col.find_one({"id": current_user["id"]})
        if not user or not user.get("enrolled_courses"):
            return None
        
        last_course_id = user["enrolled_courses"][-1]
        courses_col = get_collection('courses')
        course = await courses_col.find_one({"id": last_course_id})
        if course:
            course.pop('_id', None)
        return course
    except Exception:
        return None

@router.get("/recommendations")
async def get_course_recommendations(current_user: dict = Depends(get_current_user)):
    """Get personalized course recommendations"""
    try:
        courses_col = get_collection('courses')
        # Simple recommendation: return 3 latest courses
        courses = await courses_col.find({}).sort("created_at", -1).limit(3).to_list(3)
        for course in courses:
            course.pop('_id', None)
        return courses
    except Exception:
        return []

@router.post("/{course_id}/upload-file")
async def upload_course_file(
    course_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a content file for a course (Author only, DRAFT only)"""
    # 1. Vérifier l'existence et le statut du cours
    courses_col = get_collection('courses')
    course = await courses_col.find_one({"id": course_id, "author": current_user["id"]})
    
    if not course:
        raise HTTPException(status_code=404, detail="Cours non trouvé ou accès refusé")
    
    # Sécurité : On ne peut ajouter des fichiers qu'à un brouillon
    if course.get("status") == "published":
        raise HTTPException(status_code=400, detail="Impossible d'ajouter des fichiers à un cours déjà publié. Créez un nouveau brouillon.")
    
    # 2. Valider le type de fichier
    allowed_extensions = ['.pdf', '.docx', '.jpg', '.jpeg', '.png', '.webp', '.txt']
    allowed_content_types = [
        "application/pdf", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "image/jpeg", "image/png", "image/webp", "text/plain"
    ]
    
    fname = file.filename or ""
    import re
    fname = re.sub(r'[^a-zA-Z0-9._-]', '_', fname)
    file_extension = os.path.splitext(fname)[1].lower()
    
    if file_extension not in allowed_extensions and file.content_type not in allowed_content_types:
        logger.warning(f"Tentative d'upload de fichier non supporté: {fname} ({file.content_type})")
        raise HTTPException(status_code=400, detail="Type de fichier non supporté (PDF, DOCX, Images, TXT uniquement)")
    
    # 3. Sauvegarde physique
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    save_dir = os.path.join(base_dir, "uploads/course_contents")
    os.makedirs(save_dir, exist_ok=True)
    
    unique_filename = f"{course_id}_{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(save_dir, unique_filename)

    try:
        content = await file.read()
        # Max size validation (10MB)
        if len(content) > 10 * 1024 * 1024:
             raise HTTPException(status_code=400, detail="Le fichier est trop volumineux (max 10Mo)")
             
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        
        file_url = f"/uploads/course_contents/{unique_filename}"
        
        ctype = file.content_type or ""
        file_type = "pdf" if "pdf" in ctype else "doc" if "word" in ctype else "image" if "image" in ctype else "text" if "text" in ctype else "autre"
        
        file_doc = {
            "id": str(uuid.uuid4()),
            "courseId": course_id,
            "fileUrl": file_url,
            "fileType": file_type,
            "fileName": fname,
            "mimeType": ctype,
            "size": len(content),
            "originalFilename": fname,
            "createdAt": datetime.now(timezone.utc)
        }
        
        files_col = get_collection('course_files')
        await files_col.insert_one(file_doc)
        file_doc.pop('_id', None)
        
        logger.info(f"Fichier uploadé avec succès pour le cours {course_id}: {fname}")
        return file_doc
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur upload fichier cours: {e}")
        raise HTTPException(status_code=500, detail="Erreur lors de l'enregistrement du fichier")

@router.post("/upload-cover")
async def upload_course_cover(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a cover image for a course"""
    # Valider le type de fichier
    allowed_types = ["image/jpeg", "image/png", "image/webp"]
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail="Seuls les fichiers JPG, PNG et WEBP sont acceptés"
        )
    
    # Valider la taille (max 5MB)
    MAX_SIZE = 5 * 1024 * 1024
    content = await file.read()
    if len(content) > MAX_SIZE:
        raise HTTPException(
            status_code=400,
            detail="L'image est trop volumineuse (max 5Mo)"
        )
    await file.seek(0)

    # Créer le nom de fichier unique
    filename = file.filename or "cover.jpg"
    file_extension = os.path.splitext(filename)[1]
    if not file_extension:
        file_extension = ".jpg" # fallback
    
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    
    # On utilise un chemin absolu pour garantir l'emplacement
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) # racine du dossier backend
    save_path = os.path.join(base_dir, "uploads/courses", unique_filename)
    os.makedirs(os.path.dirname(save_path), exist_ok=True)

    try:
        with open(save_path, "wb") as buffer:
            buffer.write(content)
        
        # URL relative
        url = f"/uploads/courses/{unique_filename}"
        logger.info(f"Image de couverture sauvegardée: {save_path} -> URL: {url}")
        return {"url": url}
    except Exception as e:
        logger.error(f"Erreur upload image: {e}")
        raise HTTPException(status_code=500, detail="Erreur lors de l'enregistrement de l'image")

def get_collection(name: str):
    """Get collection from db module"""
    from db import get_collection as gc
    return gc(name)

@router.post("/seed")
async def seed_courses(current_user: dict = Depends(get_current_user)):
    """Seed sample courses"""
    try:
        courses_col = get_collection('courses')
        count = await courses_col.count_documents({"author": current_user["id"]})
        if count > 0:
            return {"message": "Database already seeded for this user"}
            
        sample_courses = [
            {
                "id": str(uuid.uuid4()),
                "title": "Introduction au Droit OHADA",
                "description": "Comprendre les bases du droit des affaires en Afrique.",
                "coverImage": None,
                "author": current_user["id"],
                "category": "comptabilité",
                "level": "débutant",
                "chapters": [],
                "aiStatus": "idle",
                "aiSummary": None,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            },
            {
                "id": str(uuid.uuid4()),
                "title": "Comptabilité Générale",
                "description": "Maîtriser le plan comptable SYSCOHADA.",
                "coverImage": None,
                "author": current_user["id"],
                "category": "comptabilité",
                "level": "intermédiaire",
                "chapters": [],
                "aiStatus": "idle",
                "aiSummary": None,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            }
        ]
        await courses_col.insert_many(sample_courses)
        return {"message": f"Seeded {len(sample_courses)} courses"}
    except Exception as e:
        logger.error(f"Seed error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/my-drafts")
@router.get("/drafts")
async def get_my_drafts(current_user: dict = Depends(get_current_user)):
    """Get author's draft courses specifically"""
    try:
        courses_col = get_collection('courses')
        query = {"author": current_user["id"], "status": "draft"}
        courses = await courses_col.find(query).sort("updated_at", -1).to_list(100)
        for course in courses:
            course.pop('_id', None)
        return courses
    except Exception as e:
        logger.error(f"Error fetching drafts: {e}")
        return []

@router.get("")
async def get_courses(subject: Optional[str] = None, skip: int = 0, limit: int = 50, current_user: dict = Depends(get_current_user_optional)):
    """Get all courses (filtered by current user or public published ones)"""
    try:
        courses_col = get_collection('courses')
        
        # Filtre de base : cours publiés OU cours dont l'utilisateur est l'auteur
        query = {"$or": [{"status": "published"}]}
        
        if current_user:
            query["$or"].append({"author": current_user["id"]})
            
        if subject:
            query["category"] = subject
            
        courses = await courses_col.find(query).sort("created_at", -1).skip(skip).to_list(limit)
        
        for course in courses:
            course.pop('_id', None)
            
        return courses
    except Exception as e:
        logger.error(f"Error fetching courses: {e}")
        return []

@router.get("/{course_id}")
async def get_course(course_id: str, current_user: dict = Depends(get_current_user)):
    """Get course details (with safety checks)"""
    try:
        courses_col = get_collection('courses')
        course = await courses_col.find_one({"id": course_id})
        
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
            
        is_author = course.get("author") == current_user["id"]
        is_published = course.get("status") == "published"
        
        if not is_author and not is_published:
            raise HTTPException(status_code=403, detail="Cours introuvable ou vous n'êtes pas autorisé à le voir.")
        
        course.pop('_id', None)
        return course
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching course {course_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("")
async def create_course(course_data: dict, current_user: dict = Depends(get_current_user)):
    """Create a new course"""
    try:
        courses_col = get_collection('courses')
        
        # Valider les données
        if not course_data.get('title'):
            raise HTTPException(status_code=400, detail="Title is required")
        
        # Créer le cours
        now = datetime.now(timezone.utc)
        course = {
            "id": course_data.get('id', str(uuid.uuid4())),
            "title": course_data['title'],
            "description": course_data.get('description'),
            "coverImage": course_data.get('coverImage'),
            "author": current_user["id"],
            "category": course_data.get('category', 'informatique'),
            "level": course_data.get('level', 'débutant'),
            "status": "draft",
            "pages": [],
            "chapters": course_data.get('chapters', []),
            "enrolled_users": [],
            "aiStatus": "idle",
            "aiSummary": None,
            "created_at": now,
            "updated_at": now
        }
        
        await courses_col.insert_one(course)
        course.pop('_id', None)
        
        logger.info(f"Course created successfully in DB with ID: {course['id']}")
        return course
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating course: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/{course_id}/enroll")
async def enroll_course(user_id: str, course_id: str):
    """Enroll in a course"""
    try:
        courses_col = get_collection('courses')
        users_col = get_collection('users')
        
        course = await courses_col.find_one({"id": course_id})
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
        
        await courses_col.update_one(
            {"id": course_id},
            {"$addToSet": {"enrolled_users": user_id}}
        )
        await users_col.update_one(
            {"id": user_id},
            {"$addToSet": {"enrolled_courses": course_id}}
        )
        return {"status": "enrolled"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error enrolling user {user_id} to course {course_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/{course_id}/pages")
async def update_course_pages(
    course_id: str,
    pages_data: list,
    current_user: dict = Depends(get_current_user)
):
    """Update all pages of a course (order, title, etc) with versioning"""
    try:
        courses_col = get_collection('courses')
        course = await courses_col.find_one({"id": course_id, "author": current_user["id"]})
        if not course:
            raise HTTPException(status_code=404, detail="Cours non trouvé")

        # Versioning : Si le cours est publié, on crée un brouillon pour les modifications
        if course.get("status") == "published":
            draft_id = f"draft_{uuid.uuid4()}"
            new_draft = {
                **course,
                "id": draft_id,
                "parent_course_id": course_id,
                "status": "draft",
                "pages": pages_data,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            }
            new_draft.pop('_id', None)
            await courses_col.insert_one(new_draft)
            return {"status": "success", "new_draft_id": draft_id, "message": "Nouveau brouillon créé pour les modifications."}

        # Sinon, mise à jour normale du brouillon existant
        await courses_col.update_one(
            {"id": course_id},
            {"$set": {"pages": pages_data, "updated_at": datetime.now(timezone.utc)}}
        )
        return {"status": "success"}
    except Exception as e:
        logger.error(f"Error updating pages: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/{course_id}/import-file/{file_id}")
async def import_file_to_pages(
    course_id: str,
    file_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Normalize file content and create course pages with technical diagnostic"""
    try:
        logger.info(f"Démarrage diagnostic importation: course={course_id}, file={file_id}")
        courses_col = get_collection('courses')
        course = await courses_col.find_one({"id": course_id, "author": current_user["id"]})
        if not course:
            logger.error(f"Diagnostic échec: Cours {course_id} non trouvé")
            raise HTTPException(status_code=404, detail="Cours non trouvé")

        files_col = get_collection('course_files')
        file_record = await files_col.find_one({"id": file_id, "courseId": course_id})
        if not file_record:
            logger.error(f"Diagnostic échec: Fichier {file_id} non trouvé en BD")
            raise HTTPException(status_code=404, detail="Fichier non trouvé")

        # Chemin absolu du fichier
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        # Correction CRITIQUE: Le chemin stocké en BD peut être absolu ou relatif
        file_url = file_record["fileUrl"]
        
        # Si c'est une URL relative commençant par /uploads
        if file_url.startswith("/uploads/"):
            file_path = os.path.join(base_dir, file_url.lstrip("/"))
        elif "course_contents/" in file_url:
            # Si le chemin contient déjà le dossier, on essaie de reconstruire proprement
            filename = os.path.basename(file_url)
            file_path = os.path.join(base_dir, "uploads/course_contents", filename)
        else:
            # Fallback sur la logique précédente si besoin
            file_path = os.path.join(base_dir, file_url.lstrip("/"))

        if not os.path.exists(file_path):
            # Tentative de recherche dans les dossiers connus
            search_paths = [
                os.path.join(base_dir, "uploads/course_contents", os.path.basename(file_url)),
                os.path.join(base_dir, "backend/uploads/course_contents", os.path.basename(file_url))
            ]
            for p in search_paths:
                if os.path.exists(p):
                    file_path = p
                    break
            
            if not os.path.exists(file_path):
                logger.error(f"Diagnostic échec: Fichier physique absent. Tenté: {file_path}")
                raise HTTPException(status_code=404, detail=f"Fichier physique introuvable (Cible: {os.path.basename(file_url)})")

        file_size = os.path.getsize(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        mime = file_record.get("mimeType", "inconnu")
        logger.info(f"Analyse document: Type={ext}, MIME={mime}, Taille={file_size} octets")

        normalized_text = ""

        # 1. Extraction contrôlée
        if ext == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    total_pages = len(reader.pages)
                    logger.info(f"PDF détecté: {total_pages} pages physiques")
                    
                    extracted_parts = []
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text() or ""
                        if text.strip():
                            extracted_parts.append(text)
                    
                    normalized_text = "\n\n--- PAGE_BREAK ---\n\n".join(extracted_parts)
                    
                    if not normalized_text.strip():
                        logger.warning("Diagnostic: PDF semble être une image (aucun texte extrait)")
                        raise HTTPException(
                            status_code=400, 
                            detail="Ce PDF semble être une image scannée. L'extraction de texte n'est pas possible pour le moment."
                        )
            except Exception as e:
                logger.error(f"Erreur technique PDF: {str(e)}")
                raise HTTPException(status_code=400, detail=f"Erreur technique lors de la lecture du PDF: {str(e)}")

        elif ext == ".docx":
            try:
                doc = docx.Document(file_path)
                logger.info(f"DOCX détecté: {len(doc.paragraphs)} paragraphes")
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                normalized_text = "\n\n".join(paragraphs)
            except Exception as e:
                logger.error(f"Erreur technique DOCX: {str(e)}")
                raise HTTPException(status_code=400, detail=f"Erreur technique lors de la lecture du DOCX: {str(e)}")

        elif ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    normalized_text = f.read()
                logger.info("TXT détecté et lu")
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    normalized_text = f.read()
                logger.info("TXT détecté (encodage latin-1)")

        # 2. Normalisation et découpage final
        if not normalized_text.strip():
            logger.error("Diagnostic final: Contenu vide après normalisation")
            raise HTTPException(status_code=400, detail="Le document ne contient aucun texte exploitable.")

        # Découpage intelligent du texte normalisé
        final_pages = []
        # Si c'était un PDF, on respecte les sauts de page extraits
        if "--- PAGE_BREAK ---" in normalized_text:
            sections = normalized_text.split("--- PAGE_BREAK ---")
        else:
            # Sinon découpage par blocs de ~2000 caractères pour le confort de lecture
            import re
            sections = re.findall(r'.{1,2000}(?:\s|$)', normalized_text, re.DOTALL)

        for i, content in enumerate(sections):
            if content.strip():
                final_pages.append({
                    "id": str(uuid.uuid4()),
                    "title": f"Page {i+1}",
                    "content": content.strip(),
                    "order": i
                })

        logger.info(f"Diagnostic Succès: {len(final_pages)} pages générées")
        
        await courses_col.update_one(
            {"id": course_id},
            {"$set": {"pages": final_pages, "updated_at": datetime.now(timezone.utc)}}
        )

        return {"status": "success", "pages_count": len(final_pages), "diagnostic": "normalisation_complete"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ERREUR CRITIQUE IMPORTATION: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur technique imprévue: {str(e)}")

@router.post("/{course_id}/publish")
@router.patch("/{course_id}/publish")
async def publish_course(
    course_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Publish a course (DRAFT -> PUBLISHED)"""
    try:
        courses_col = get_collection('courses')
        course = await courses_col.find_one({"id": course_id, "author": current_user["id"]})
        
        if not course:
            raise HTTPException(status_code=404, detail="Cours non trouvé")
            
        # On autorise la publication même si le contenu est minimal pour les tests
        # Mais on garde une trace
        if not course.get("pages") or len(course.get("pages")) == 0:
            logger.warning(f"Cours {course_id} publié sans pages de contenu")

        await courses_col.update_one(
            {"id": course_id},
            {"$set": {"status": "published", "updated_at": datetime.now(timezone.utc)}}
        )
        return {"status": "published"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error publishing course: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/{course_id}")
async def delete_course(
    course_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a course and its associated files"""
    try:
        courses_col = get_collection('courses')
        course = await courses_col.find_one({"id": course_id, "author": current_user["id"]})
        
        if not course:
            raise HTTPException(status_code=404, detail="Cours non trouvé ou accès refusé")

        # 1. Supprimer les fichiers associés
        files_col = get_collection('course_files')
        files = await files_col.find({"courseId": course_id}).to_list(None)
        
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        for file in files:
            try:
                file_url = file.get("fileUrl", "")
                if file_url.startswith("/uploads/"):
                    file_path = os.path.join(base_dir, file_url.lstrip("/"))
                    if os.path.exists(file_path):
                        os.remove(file_path)
            except Exception as e:
                logger.error(f"Erreur suppression fichier physique: {e}")

        # 2. Supprimer les enregistrements en BD
        await files_col.delete_many({"courseId": course_id})
        await courses_col.delete_one({"id": course_id})
        
        # 3. Mettre à jour les utilisateurs qui auraient ce cours dans leurs favoris/enrolled
        users_col = get_collection('users')
        await users_col.update_many(
            {},
            {"$pull": {"enrolled_courses": course_id}}
        )

        return {"status": "success", "message": "Cours supprimé avec succès"}
    except Exception as e:
        logger.error(f"Error deleting course: {e}")
        raise HTTPException(status_code=500, detail=str(e))
